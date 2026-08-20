from pathlib import Path

from docx import Document

from app.models.export import ReportTemplate, ReportTemplateMapping
from app.models.generation import DocumentTemplate, SectionDraft, SectionDraftVersion, SectionGenerationConfig, TemplateSection
from app.models.review import ProfessionalReviewRun, QualityScoreResult
from app.models.structured_data import RawMaterial
from app.models.workflow import DocumentSectionInstance
from app.services import export_service, storage


def _auth(client):
    client.post("/api/v1/auth/register", json={"username": "exportuser", "email": "export@example.com", "password": "secret123"})
    token = client.post("/api/v1/auth/login", json={"username": "exportuser", "password": "secret123"}).json()["data"]["access_token"]
    return {"Authorization": f"Bearer {token}"}


def _template(db, user_id):
    template = DocumentTemplate(name="TEST 导出内容模板", code="test_export_document", document_type="emergency_response", status="active")
    db.add(template); db.flush()
    section = TemplateSection(template_id=template.id, section_code="1", title="原辅材料情况", level=1, sort_order=1, generation_mode="facts_only", required=True)
    db.add(section); db.flush(); db.add(SectionGenerationConfig(section_id=section.id, prompt_template="section_base_v1", required_fields=[], prompt_version="export-v1")); db.commit()
    docx = Document(); docx.add_paragraph("{{ report_title }}", style="Title"); docx.add_paragraph("{{ company_name }}"); docx.add_paragraph("{{ report_body }}")
    rel_path = f"report_templates/{user_id}/test-export-template.docx"; path = storage.get_storage().resolve_path(rel_path); path.parent.mkdir(parents=True, exist_ok=True); docx.save(path)
    report_template = ReportTemplate(name="EnvAI TEST Report Template", code="test_export_word", document_type="emergency_response", version="v1", original_file_name="test.docx", storage_path=rel_path, sha256=export_service._sha256(path), file_size=path.stat().st_size, created_by=user_id)
    db.add(report_template); db.flush(); db.add(ReportTemplateMapping(report_template_id=report_template.id, document_template_id=template.id, section_mappings={})); db.commit()
    return template, section, report_template


def test_snapshot_docx_pdf_and_access(client, db):
    headers = _auth(client)
    user_id = client.get("/api/v1/auth/me", headers=headers).json()["data"]["id"]
    template, section, report_template = _template(db, user_id)
    project = client.post("/api/v1/projects", headers=headers, json={"name": "导出测试项目"}).json()["data"]
    client.put(f"/api/v1/projects/{project['id']}/profile", headers=headers, json={"company_name": "江苏测试环保科技有限公司", "project_address": "苏州市测试园区"})
    instance = client.post(f"/api/v1/projects/{project['id']}/document-instances", headers=headers, json={"template_id": template.id, "title": "测试应急预案"}).json()["data"]
    instance_id = instance["id"]
    snapshot_section = db.query(DocumentSectionInstance).filter_by(document_instance_id=instance_id, template_section_id=section.id).one()
    draft = SectionDraft(project_id=project["id"], document_instance_id=instance_id, template_id=template.id, section_id=section.id, created_by=user_id, content="甲苯为本项目原辅材料。", status="approved", version=1)
    db.add(draft); db.flush(); version = SectionDraftVersion(draft_id=draft.id, version=1, content=draft.content, status="approved", saved_by=user_id); db.add(version); db.flush()
    snapshot_section.current_draft_id = draft.id; snapshot_section.approved_version_id = version.id; snapshot_section.status = "approved"
    db.add(RawMaterial(project_id=project["id"], name="甲苯", annual_usage="35", annual_usage_unit="t/a", max_storage="0.035", storage_unit="t", storage_location="甲类仓库", verification_status="user_verified"))
    run = ProfessionalReviewRun(document_instance_id=instance_id, status="completed", review_mode="rules_only", started_by=user_id); db.add(run); db.flush()
    db.add(QualityScoreResult(document_instance_id=instance_id, review_run_id=run.id, overall_score=100, data_integrity_score=100, citation_score=100, coverage_score=100, completeness_score=100, consistency_score=100, critical_issue_count=0, major_issue_count=0, quality_passed=True)); db.commit()
    preflight = client.post(f"/api/v1/document-instances/{instance_id}/export-preflight?report_template_id={report_template.id}", headers=headers).json()["data"]
    assert preflight["ready"] is True
    frozen = client.post(f"/api/v1/document-instances/{instance_id}/snapshots", headers=headers, json={}).json()["data"]
    client.put(f"/api/v1/section-drafts/{draft.id}", headers=headers, json={"content": "正文后来已修改。"})
    stored = client.get(f"/api/v1/report-snapshots/{frozen['id']}", headers=headers).json()["data"]
    assert stored["content_hash"] == frozen["content_hash"]
    job = client.post(f"/api/v1/report-snapshots/{frozen['id']}/exports", headers=headers, json={"formats": ["docx", "pdf"], "report_template_id": report_template.id}).json()["data"]
    assert job["docx_status"] == "completed"
    assert job["status"] in {"completed", "partial"}
    artifacts = client.get(f"/api/v1/report-export-jobs/{job['id']}/artifacts", headers=headers).json()["data"]
    assert any(item["format"] == "docx" and len(item["sha256"]) == 64 for item in artifacts)
    docx_artifact = next(item for item in artifacts if item["format"] == "docx")
    response = client.get(f"/api/v1/export-artifacts/{docx_artifact['id']}/download", headers=headers)
    assert response.status_code == 200 and response.content[:2] == b"PK"
    client.post("/api/v1/auth/register", json={"username": "exportother", "email": "exportother@example.com", "password": "secret123"})
    other = client.post("/api/v1/auth/login", json={"username": "exportother", "password": "secret123"}).json()["data"]["access_token"]
    assert client.get(f"/api/v1/export-artifacts/{docx_artifact['id']}/download", headers={"Authorization": f"Bearer {other}"}).json()["code"] == 404
