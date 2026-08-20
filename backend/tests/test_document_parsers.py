import pymupdf as fitz
import openpyxl
from docx import Document
from PIL import Image

from app.services.document_parser import ParserRegistry
from app.services.document_parser.docx_parser import DOCXParser
from app.services.document_parser.excel_parser import ExcelParser
from app.services.document_parser.image_parser import ImageParser
from app.services.document_parser.normalizer import TextNormalizer
from app.services.document_parser.pdf_parser import PDFParser


def _make_pdf(path):
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Environmental project 35 t/a")
    document.save(path)
    document.close()


def test_pdf_parser_pages_and_text(tmp_path):
    path = tmp_path / "sample.pdf"
    _make_pdf(path)
    result = PDFParser().parse(path)
    assert len(result.pages) == 1
    assert "35 t/a" in result.plain_text
    assert result.metadata["page_count"] == 1


def test_pdf_parser_detects_scanned_document(tmp_path):
    path = tmp_path / "scanned.pdf"
    document = fitz.open()
    document.new_page()
    document.save(path)
    document.close()
    result = PDFParser().parse(path)
    assert result.metadata["requires_ocr"] is True
    assert "possible_scanned_pdf" in result.warnings


def test_docx_parser_paragraphs_and_tables(tmp_path):
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("企业基本情况", level=1)
    document.add_paragraph("测试环保科技有限公司，年产量 5000 t/a")
    table = document.add_table(rows=2, cols=3)
    for cell, value in zip(table.rows[0].cells, ["原辅材料名称", "年使用量", "单位"]):
        cell.text = value
    for cell, value in zip(table.rows[1].cells, ["甲苯", "35", "t/a"]):
        cell.text = value
    document.save(path)

    result = DOCXParser().parse(path)
    assert result.paragraphs[0]["style"].lower().startswith("heading")
    assert result.tables[0]["rows"][1] == ["甲苯", "35", "t/a"]
    assert "5000 t/a" in result.plain_text


def test_excel_parser_sheets_values_and_formula(tmp_path):
    path = tmp_path / "sample.xlsx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "原辅材料"
    sheet.append(["名称", "年用量", "单位"])
    sheet.append(["甲苯", 35, "t/a"])
    sheet.append(["乙醇", 10, "t/a"])
    sheet["D1"] = "公式"
    sheet["D2"] = "=B2*2"
    workbook.create_sheet("企业信息")["A1"] = "测试环保科技有限公司"
    workbook.save(path)

    result = ExcelParser().parse(path)
    materials = next(item for item in result.sheets if item["name"] == "原辅材料")
    assert materials["rows"][1][:3] == ["甲苯", 35, "t/a"]
    assert materials["rows"][1][3] == "=B2*2"
    assert "测试环保科技有限公司" in result.plain_text


def test_image_parser_metadata(tmp_path):
    path = tmp_path / "sample.png"
    Image.new("RGB", (32, 16), "white").save(path)
    result = ImageParser().parse(path)
    assert result.metadata["width"] == 32
    assert result.metadata["height"] == 16
    assert "ocr_not_configured" in result.warnings


def test_normalizer_preserves_professional_values():
    text = "\x00 35  t/a\r\n\r\n\r\n1.5mg/m³\nGB 16297-1996"
    assert TextNormalizer.normalize(text) == "35 t/a\n\n1.5mg/m³\nGB 16297-1996"


def test_parser_registry_matches_supported_extensions():
    registry = ParserRegistry()
    assert registry.get_parser(".pdf").name == "pymupdf"
    assert registry.get_parser(".docx").name == "python-docx"
    assert registry.get_parser(".xlsx").name == "openpyxl/xlrd"
    assert registry.get_parser(".xls").name == "openpyxl/xlrd"
    assert registry.get_parser(".png").name == "pillow"
    assert registry.get_parser(".jpg").name == "pillow"
    assert registry.get_parser(".jpeg").name == "pillow"
