from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from app.core.config import settings

from .base import BaseDocumentParser, DocumentParseError, ParsedDocumentResult
from .normalizer import TextNormalizer
from .utils import ensure_safe_zip, ensure_text_limit


class DOCXParser(BaseDocumentParser):
    name = "python-docx"
    extensions = {".docx"}

    def parse(self, file_path: Path) -> ParsedDocumentResult:
        ensure_safe_zip(file_path)
        try:
            document = Document(file_path)
        except (PackageNotFoundError, ValueError, KeyError) as exc:
            raise DocumentParseError("DOCX 文件无效或已损坏") from exc

        if len(document.paragraphs) > settings.max_docx_paragraphs:
            raise DocumentParseError(f"DOCX 段落数超过 {settings.max_docx_paragraphs} 条限制")
        if len(document.tables) > settings.max_table_count:
            raise DocumentParseError(f"DOCX 表格数超过 {settings.max_table_count} 个限制")

        paragraphs = []
        text_parts = []
        for index, paragraph in enumerate(document.paragraphs):
            text = TextNormalizer.normalize(paragraph.text)
            style = paragraph.style.name if paragraph.style else ""
            paragraphs.append(
                {
                    "index": index,
                    "text": text,
                    "style": style,
                    "is_heading": style.lower().startswith("heading") or style.startswith("标题"),
                }
            )
            if text:
                text_parts.append(text)

        tables = []
        for table_index, table in enumerate(document.tables):
            if len(table.rows) > settings.max_table_rows:
                raise DocumentParseError(f"DOCX 第 {table_index + 1} 个表格行数超过限制")
            rows = [
                [TextNormalizer.normalize(cell.text) for cell in row.cells]
                for row in table.rows
            ]
            tables.append({"index": table_index, "rows": rows})
            text_parts.extend(" | ".join(cell for cell in row if cell) for row in rows)

        return ParsedDocumentResult(
            parser_name=self.name,
            parser_version=self.version,
            plain_text=ensure_text_limit(TextNormalizer.normalize("\n".join(text_parts))),
            paragraphs=paragraphs,
            tables=tables,
            metadata={"paragraph_count": len(paragraphs), "table_count": len(tables)},
        )
