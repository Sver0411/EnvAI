from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import subprocess
import tempfile
import uuid
import zipfile
from datetime import datetime, timezone
from decimal import Decimal
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docxtpl import DocxTemplate
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.core.config import settings
from app.core.exceptions import NotFoundError, ValidationError
from app.models.company_profile import CompanyProfile
from app.models.export import ExportArtifact, ReportExportJob, ReportFigure, ReportSnapshot, ReportTemplate, ReportTemplateMapping
from app.models.generation import DocumentInstance, DocumentTemplate, SectionCitation, SectionDraftVersion
from app.models.knowledge import KnowledgeChunk, KnowledgeDocument
from app.models.project_file import ProjectFile
from app.models.review import ProfessionalReviewRun
from app.models.structured_data import Product, ProductionEquipment, RawMaterial
from app.models.user import User
from app.models.workflow import DocumentSectionInstance
from app.services import generation_service, storage, workflow_service
from app.services.review_service import quality_gate
from app.services.authorization import current_organization
from app.services import tenant_service


EXPORTER_VERSION = "docx_exporter_v1"
INTERNAL_MARKER = re.compile(r"\[\[(?:KB|PROJECT|FACT):[^\]]+\]\]")
PLACEHOLDER = re.compile(r"\{\{[^{}]+\}\}")
ALLOWED_TEMPLATE_PLACEHOLDERS = {"document_title", "report_title", "company_name", "project_address", "report_date", "report_type", "snapshot_version", "draft_mark", "report_body"}


def _value(value: Any) -> Any:
    if isinstance(value, Decimal):
        return format(value.normalize(), "f").rstrip("0").rstrip(".") or "0"
    if isinstance(value, datetime):
        return value.isoformat()
    return value


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def safe_filename(value: str, suffix: str) -> str:
    safe = re.sub(r"[\\/:*?\"<>|\x00-\x1f]+", "_", value).strip(" ._")
    safe = re.sub(r"\s+", " ", safe)[:160] or "EnvAI_Report"
    return f"{safe}{suffix}"


def _owned_instance(db: Session, instance_id: int, user: User) -> DocumentInstance:
    return generation_service.get_instance(db, instance_id, user)


def get_template(db: Session, template_id: int, user: User) -> ReportTemplate:
    org = current_organization(db, user)
    template = db.scalar(select(ReportTemplate).where(ReportTemplate.id == template_id, ReportTemplate.status != "archived", (ReportTemplate.created_by == user.id) | (ReportTemplate.organization_id == org.id)))
    if template is None:
        raise NotFoundError("报告模板不存在")
    return template


def list_templates(db: Session, user: User) -> list[ReportTemplate]:
    org = current_organization(db, user)
    return list(db.scalars(select(ReportTemplate).where(ReportTemplate.status != "archived", (ReportTemplate.created_by == user.id) | (ReportTemplate.organization_id == org.id)).order_by(ReportTemplate.created_at.desc())))


def validate_template_file(path: Path) -> list[str]:
    warnings: list[str] = []
    if path.suffix.lower() != ".docx":
        raise ValidationError("报告模板只允许 .docx，不允许宏模板 .docm")
    try:
        with zipfile.ZipFile(path) as archive:
            infos = archive.infolist()
            if len(infos) > 10_000:
                raise ValidationError("模板压缩包条目过多")
            if sum(item.file_size for item in infos) > settings.max_archive_uncompressed_size_mb * 1024 * 1024:
                raise ValidationError("模板解压后体积超限")
            names = set(archive.namelist())
            if "word/document.xml" not in names or "[Content_Types].xml" not in names:
                raise ValidationError("不是有效的 Word DOCX 模板")
            if any(name.lower().endswith("vbaproject.bin") for name in names):
                raise ValidationError("模板包含宏，禁止上传")
    except zipfile.BadZipFile as exc:
        raise ValidationError("报告模板不是有效 DOCX 文件") from exc
    document = Document(path)
    styles = {style.name for style in document.styles}
    missing = [name for name in ("Normal", "Heading 1", "Heading 2", "Heading 3") if name not in styles]
    if missing:
        warnings.append(f"模板缺少建议样式：{', '.join(missing)}")
    return warnings


def _validate_template_placeholders(path: Path) -> None:
    with zipfile.ZipFile(path) as archive:
        xml = "\n".join(archive.read(name).decode("utf-8", errors="ignore") for name in archive.namelist() if name.startswith("word/") and name.endswith(".xml"))
    placeholders = {item.strip() for item in re.findall(r"\{\{\s*([^{}]+?)\s*\}\}", xml)}
    unknown = sorted(placeholders - ALLOWED_TEMPLATE_PLACEHOLDERS)
    if unknown:
        raise ValidationError(f"Word 模板包含未支持的占位符：{', '.join(unknown)}")


def create_template(db: Session, *, name: str, code: str, document_type: str, version: str, document_template_id: int, original_file_name: str, stored_path: str, file_size: int, user: User, section_mappings: dict[str, Any] | None = None) -> ReportTemplate:
    path = storage.get_storage().resolve_path(stored_path)
    validate_template_file(path)
    if db.scalar(select(DocumentTemplate).where(DocumentTemplate.id == document_template_id)) is None:
        raise NotFoundError("内容模板不存在")
    if db.scalar(select(ReportTemplate).where(ReportTemplate.code == code, ReportTemplate.version == version)):
        raise ValidationError("报告模板编码和版本已存在")
    organization = current_organization(db, user)
    template = ReportTemplate(name=name, code=code, document_type=document_type, version=version, original_file_name=safe_filename(original_file_name, "")[:255], storage_path=stored_path, sha256=_sha256(path), file_size=file_size, created_by=user.id, organization_id=organization.id)
    db.add(template); db.flush()
    db.add(ReportTemplateMapping(report_template_id=template.id, document_template_id=document_template_id, section_mappings=section_mappings or {}))
    db.commit(); db.refresh(template)
    return template


def ensure_default_report_template(db: Session, document_template_id: int, user: User) -> ReportTemplate:
    """Create the local system DOCX layout on first use.

    A customer may upload a house style later, but exporting a valid report
    must not depend on that upload.  The generated file is still represented
    by the normal template/mapping tables so existing snapshot and export
    audit paths remain unchanged.
    """
    organization = current_organization(db, user)
    content_template = db.get(DocumentTemplate, document_template_id)
    if content_template is None:
        raise NotFoundError("内容模板不存在")
    visible = db.scalar(
        select(ReportTemplate)
        .join(ReportTemplateMapping)
        .where(
            ReportTemplateMapping.document_template_id == document_template_id,
            ReportTemplate.status == "active",
            (ReportTemplate.created_by == user.id) | (ReportTemplate.organization_id == organization.id),
        )
        .order_by(ReportTemplate.id.desc())
    )
    if visible is not None:
        return visible

    code = f"envai_system_default_{document_template_id}"
    template = db.scalar(select(ReportTemplate).where(ReportTemplate.code == code, ReportTemplate.version == "v1"))
    rel_path = f"report_templates/system/{code}.docx"
    path = storage.get_storage().resolve_path(rel_path)
    if template is None:
        path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        title = document.add_paragraph("{{ report_title }}", style="Title")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            _set_east_asia(run)
        for label, key in (("企业名称", "company_name"), ("项目地址", "project_address"), ("报告日期", "report_date"), ("版本", "snapshot_version")):
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{label}：").bold = True
            paragraph.add_run(f"{{{{ {key} }}}}")
            for run in paragraph.runs:
                _set_east_asia(run)
        document.add_page_break()
        document.save(path)
        template = ReportTemplate(
            name=f"系统默认 Word 模板（{content_template.name}）",
            code=code,
            document_type=content_template.document_type,
            version="v1",
            status="active",
            original_file_name="envai-system-default.docx",
            storage_path=rel_path,
            sha256=_sha256(path),
            file_size=path.stat().st_size,
            engine="docxtpl",
            created_by=user.id,
            organization_id=organization.id,
        )
        db.add(template)
        db.flush()
        db.add(ReportTemplateMapping(report_template_id=template.id, document_template_id=document_template_id, section_mappings={}))
        db.commit()
        db.refresh(template)
    elif not path.is_file():
        raise ValidationError("系统默认 Word 模板文件不存在，请重新初始化本地存储")
    return template


def _profile_snapshot(profile: CompanyProfile | None) -> dict[str, Any]:
    if profile is None:
        return {}
    return {column.name: _value(getattr(profile, column.name)) for column in profile.__table__.columns if column.name not in {"id", "project_id", "created_at", "updated_at"}}


def _citation_snapshot(db: Session, draft_id: int) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for citation in db.scalars(select(SectionCitation).where(SectionCitation.section_draft_id == draft_id)):
        item: dict[str, Any] = {"source_type": citation.source_type, "source_id": citation.source_id, "claim_text": citation.claim_text}
        if citation.source_type == "knowledge_chunk":
            chunk = db.get(KnowledgeChunk, citation.source_id)
            document = db.get(KnowledgeDocument, chunk.knowledge_document_id) if chunk else None
            item["document"] = {"id": document.id, "title": document.title, "document_number": document.document_number, "version": document.version, "status": document.status, "source_authority": document.source_authority} if document else None
        rows.append(item)
    return rows


def _table_rows(db: Session, instance: DocumentInstance) -> dict[str, list[dict[str, str]]]:
    def raw(items: list[Any], fields: list[str]) -> list[dict[str, str]]:
        return [{field: str(_value(getattr(item, field)) or "") for field in fields} for item in items]
    return {
        "raw_materials": raw(list(db.scalars(select(RawMaterial).where(RawMaterial.project_id == instance.project_id).order_by(RawMaterial.id))), ["name", "annual_usage", "annual_usage_unit", "max_storage", "storage_unit", "storage_location"]),
        "equipment": raw(list(db.scalars(select(ProductionEquipment).where(ProductionEquipment.project_id == instance.project_id).order_by(ProductionEquipment.id))), ["name", "model", "quantity", "unit", "power", "power_unit", "location"]),
        "products": raw(list(db.scalars(select(Product).where(Product.project_id == instance.project_id).order_by(Product.id))), ["name", "annual_capacity", "unit", "specification"]),
    }


def export_preflight(db: Session, instance_id: int, user: User, report_template_id: int | None = None, *, draft: bool = False) -> dict[str, Any]:
    instance = _owned_instance(db, instance_id, user)
    from app.services.authorization import require_project_permission
    require_project_permission(db, user, instance.project, "documents.export" if not draft else "documents.read")
    warnings: list[str] = []
    blocking: list[str] = []
    template: ReportTemplate | None = None
    if report_template_id:
        template = get_template(db, report_template_id, user)
        mapped = db.scalar(select(ReportTemplateMapping).where(ReportTemplateMapping.report_template_id == template.id, ReportTemplateMapping.document_template_id == instance.template_id))
        if mapped is None:
            blocking.append("所选 Word 模板未绑定当前内容模板")
    else:
        organization = current_organization(db, user)
        template = db.scalar(select(ReportTemplate).join(ReportTemplateMapping).where(ReportTemplate.status == "active", ReportTemplateMapping.document_template_id == instance.template_id, (ReportTemplate.created_by == user.id) | (ReportTemplate.organization_id == organization.id)).order_by(ReportTemplate.id.desc()))
    if template is None:
        template = ensure_default_report_template(db, instance.template_id, user)
        warnings.append("未选择自定义 Word 模板，已自动使用系统默认版式")
    profile = db.scalar(select(CompanyProfile).where(CompanyProfile.project_id == instance.project_id))
    if not instance.title.strip(): blocking.append("报告标题不能为空")
    if not profile or not profile.company_name: blocking.append("封面缺少企业名称")
    if not draft:
        readiness = workflow_service.readiness(db, instance, user)
        blocking.extend(readiness["blocking_reasons"])
        gate = quality_gate(db, instance.id, user)
        if not gate.get("passed"):
            blocking.append(gate.get("reason") or "专业质量门禁未通过")
    figures = list(db.scalars(select(ReportFigure).where(ReportFigure.document_instance_id == instance.id, ReportFigure.enabled.is_(True))))
    for figure in figures:
        project_file = db.get(ProjectFile, figure.project_file_id)
        if not project_file or not project_file.storage_path or not storage.get_storage().resolve_path(project_file.storage_path).is_file():
            blocking.append(f"图件 {figure.caption} 的源文件不存在")
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice: warnings.append("当前环境未发现 LibreOffice：仍可生成 DOCX，但 PDF 不可用")
    return {"ready": not blocking, "blocking_issues": list(dict.fromkeys(blocking)), "warnings": warnings, "selected_template": template.name if template else None, "selected_template_id": template.id if template else None, "snapshot_required": True, "pdf_available": bool(soffice)}


def create_snapshot(db: Session, instance_id: int, user: User, *, draft: bool = False) -> ReportSnapshot:
    preflight = export_preflight(db, instance_id, user, draft=draft)
    if not preflight["ready"]:
        raise ValidationError("；".join(preflight["blocking_issues"]))
    instance = _owned_instance(db, instance_id, user)
    profile = db.scalar(select(CompanyProfile).where(CompanyProfile.project_id == instance.project_id))
    sections: list[dict[str, Any]] = []
    all_citations: list[dict[str, Any]] = []
    for section in db.scalars(select(DocumentSectionInstance).where(DocumentSectionInstance.document_instance_id == instance.id).order_by(DocumentSectionInstance.sort_order)):
        version = section.approved_version
        if version is None and section.current_draft:
            # Draft snapshots explicitly preserve the editable current text;
            # formal snapshots have already passed readiness and use approvals.
            if not draft:
                raise ValidationError(f"必填章节 {section.section_code} 缺少已审核版本")
            content, draft_id, version_no = section.current_draft.content, section.current_draft.id, section.current_draft.version
        elif version is not None:
            content, draft_id, version_no = version.content, version.draft_id, version.version
        else:
            continue
        citations = _citation_snapshot(db, draft_id)
        all_citations.extend(citations)
        sections.append({"id": section.id, "section_code": section.section_code, "title": section.title, "level": section.level, "draft_version_id": version.id if version else None, "version": version_no, "content": content, "citations": citations})
    figures = []
    for figure in db.scalars(select(ReportFigure).where(ReportFigure.document_instance_id == instance.id, ReportFigure.enabled.is_(True)).order_by(ReportFigure.sort_order, ReportFigure.id)):
        project_file = db.get(ProjectFile, figure.project_file_id)
        figures.append({"id": figure.id, "section_instance_id": figure.section_instance_id, "caption": figure.caption, "width_inches": figure.width_inches, "storage_path": project_file.storage_path if project_file else None, "file_name": project_file.filename if project_file else None})
    latest_review = db.scalar(select(ProfessionalReviewRun).where(ProfessionalReviewRun.document_instance_id == instance.id).order_by(ProfessionalReviewRun.id.desc()))
    content = {"document": {"id": instance.id, "title": instance.title, "reference_date": instance.reference_date.isoformat() if instance.reference_date else None, "template_id": instance.template_id, "template_version": instance.template_version}, "company": _profile_snapshot(profile), "sections": sections, "citations": all_citations, "tables": _table_rows(db, instance), "figures": figures}
    canonical = json.dumps(content, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    number = (db.scalar(select(func.max(ReportSnapshot.snapshot_number)).where(ReportSnapshot.document_instance_id == instance.id)) or 0) + 1
    snapshot = ReportSnapshot(document_instance_id=instance.id, organization_id=instance.organization_id, snapshot_number=number, status="draft" if draft else "formal", document_title=instance.title, template_id=instance.template_id, template_version=instance.template_version, quality_review_run_id=latest_review.id if latest_review else None, snapshot_content=content, content_hash=hashlib.sha256(canonical.encode()).hexdigest(), metadata_json={"section_count": len(sections), "citation_count": len(all_citations), "figure_count": len(figures), "exporter_version": EXPORTER_VERSION}, created_by=user.id)
    db.add(snapshot); workflow_service.audit(db, user, "snapshot_created", instance, metadata={"snapshot_number": number, "draft": draft}); db.commit(); db.refresh(snapshot)
    return snapshot


def _owned_snapshot(db: Session, snapshot_id: int, user: User) -> ReportSnapshot:
    snapshot = db.get(ReportSnapshot, snapshot_id)
    if snapshot is not None:
        try: generation_service.get_instance(db, snapshot.document_instance_id, user)
        except Exception: snapshot = None
    if snapshot is None: raise NotFoundError("报告快照不存在")
    return snapshot


def list_snapshots(db: Session, instance_id: int, user: User) -> list[ReportSnapshot]:
    instance = _owned_instance(db, instance_id, user)
    return list(db.scalars(select(ReportSnapshot).where(ReportSnapshot.document_instance_id == instance.id).order_by(ReportSnapshot.snapshot_number.desc())))


def get_snapshot(db: Session, snapshot_id: int, user: User) -> ReportSnapshot:
    return _owned_snapshot(db, snapshot_id, user)


def _iter_paragraphs(document: Any):
    for paragraph in document.paragraphs: yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_paragraphs(cell)
    for section in getattr(document, "sections", []):
        yield from _iter_paragraphs(section.header)
        yield from _iter_paragraphs(section.footer)


def _set_east_asia(run) -> None:
    # A Unicode font is used for generated content; uploaded house templates
    # still retain their own styles for their fixed cover/header material.
    run.font.name = "Arial Unicode MS"
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attribute}"), "Arial Unicode MS")


def _replace_placeholders(document: Document, context: dict[str, str]) -> None:
    for paragraph in _iter_paragraphs(document):
        if not paragraph.runs: continue
        text = paragraph.text
        replaced = text
        for key, value in context.items(): replaced = replaced.replace("{{ " + key + " }}", value).replace("{{" + key + "}}", value)
        if replaced != text:
            paragraph.clear(); run = paragraph.add_run(replaced); _set_east_asia(run)


def _add_table(document: Document, headers: list[str], rows: list[dict[str, str]], keys: list[str]) -> None:
    if not rows: return
    table = document.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"; table.autofit = False
    width = 6.3 / len(headers)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]; cell.width = Inches(width); cell.text = header
        for run in cell.paragraphs[0].runs: run.bold = True; _set_east_asia(run)
    for row_data in rows:
        cells = table.add_row().cells
        for index, key in enumerate(keys):
            cells[index].width = Inches(width); cells[index].text = row_data.get(key, "")
            for run in cells[index].paragraphs[0].runs: _set_east_asia(run)
    document.add_paragraph()


def _render_markdown(document: Document, text: str) -> None:
    lines = text.replace("\r\n", "\n").split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            document.add_paragraph(); continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            paragraph = document.add_paragraph(heading.group(2), style=f"Heading {min(len(heading.group(1)) + 1, 4)}")
            for run in paragraph.runs: _set_east_asia(run)
            continue
        if re.match(r"^[-*]\s+", stripped):
            paragraph = document.add_paragraph(style="List Bullet"); content = re.sub(r"^[-*]\s+", "", stripped)
        elif re.match(r"^\d+[.)]\s+", stripped):
            paragraph = document.add_paragraph(style="List Number"); content = re.sub(r"^\d+[.)]\s+", "", stripped)
        else:
            paragraph = document.add_paragraph(); content = stripped
        parts = re.split(r"(\*\*.+?\*\*)", content)
        for part in parts:
            run = paragraph.add_run(part[2:-2] if part.startswith("**") and part.endswith("**") else part)
            run.bold = part.startswith("**") and part.endswith("**"); _set_east_asia(run)


def _replace_internal_markers(text: str, citations: list[dict[str, Any]]) -> str:
    titles = [f"《{item['document']['title']}》{('（' + item['document']['document_number'] + '）') if item.get('document') and item['document'].get('document_number') else ''}" for item in citations if item.get("document")]
    return INTERNAL_MARKER.sub(lambda match: titles[0] if titles else match.group(0), text)


def _render_docx(snapshot: ReportSnapshot, template: ReportTemplate, target: Path) -> dict[str, int]:
    source = storage.get_storage().resolve_path(template.storage_path)
    if not source.is_file(): raise ValidationError("Word 模板文件不存在")
    _validate_template_placeholders(source)
    content = snapshot.snapshot_content
    company = content.get("company", {})
    context = {"document_title": snapshot.document_title, "report_title": snapshot.document_title, "company_name": str(company.get("company_name") or ""), "project_address": str(company.get("project_address") or ""), "report_date": datetime.now().strftime("%Y年%m月%d日"), "report_type": template.document_type, "snapshot_version": f"R{snapshot.snapshot_number}", "draft_mark": "草稿 / DRAFT" if snapshot.status == "draft" else ""}
    staged = target.parent / "templated.docx"
    docxtpl = DocxTemplate(str(source)); docxtpl.render(context); docxtpl.save(staged)
    document = Document(staged); _replace_placeholders(document, context)
    if snapshot.status == "draft":
        document.add_paragraph("草稿 / DRAFT", style="Title").alignment = WD_ALIGN_PARAGRAPH.CENTER
    figures_by_section: dict[int | None, list[dict[str, Any]]] = {}
    for figure in content.get("figures", []): figures_by_section.setdefault(figure.get("section_instance_id"), []).append(figure)
    table_count = figure_count = 0
    for section in content.get("sections", []):
        heading_level = min(max(int(section.get("level", 1)), 1), 4)
        section_heading = document.add_paragraph(f"{section['section_code']} {section['title']}", style=f"Heading {heading_level}")
        for run in section_heading.runs: _set_east_asia(run)
        plain_content = generation_service.strip_markdown(str(section.get("content") or ""))
        _render_markdown(document, _replace_internal_markers(plain_content, section.get("citations") or []))
        title = section["title"]
        if "原辅材料" in title:
            _add_table(document, ["名称", "年使用量", "单位", "最大储存量", "单位", "储存位置"], content.get("tables", {}).get("raw_materials", []), ["name", "annual_usage", "annual_usage_unit", "max_storage", "storage_unit", "storage_location"]); table_count += 1
        elif "设备" in title:
            _add_table(document, ["名称", "型号", "数量", "单位", "功率", "单位", "位置"], content.get("tables", {}).get("equipment", []), ["name", "model", "quantity", "unit", "power", "power_unit", "location"]); table_count += 1
        elif "产品" in title:
            _add_table(document, ["产品名称", "年产能", "单位", "规格"], content.get("tables", {}).get("products", []), ["name", "annual_capacity", "unit", "specification"]); table_count += 1
        for figure in figures_by_section.get(section["id"], []):
            path = storage.get_storage().resolve_path(figure["storage_path"]) if figure.get("storage_path") else None
            if not path or not path.is_file(): raise ValidationError(f"图件 {figure['caption']} 不存在")
            document.add_picture(str(path), width=Inches(float(figure.get("width_inches") or 5.8)))
            caption = document.add_paragraph(f"图 {section['section_code']}-{figure_count + 1} {figure['caption']}"); caption.alignment = WD_ALIGN_PARAGRAPH.CENTER; figure_count += 1
    refs: dict[int, dict[str, Any]] = {}
    for citation in content.get("citations", []):
        document_meta = citation.get("document")
        if document_meta: refs[document_meta["id"]] = document_meta
    if refs:
        document.add_page_break(); document.add_paragraph("主要法规和标准", style="Heading 1")
        for item in sorted(refs.values(), key=lambda value: ((value.get("document_number") or ""), value["title"])):
            suffix = f"（{item['document_number']}）" if item.get("document_number") else ""
            document.add_paragraph(f"《{item['title']}》{suffix}", style="List Number")
    document.save(target)
    return {"section_count": len(content.get("sections", [])), "table_count": table_count, "figure_count": figure_count, "citation_count": len(refs)}


def validate_docx(path: Path, snapshot: ReportSnapshot) -> None:
    if not path.is_file() or path.stat().st_size < 512: raise ValidationError("DOCX 文件不存在或大小异常")
    try:
        with zipfile.ZipFile(path) as archive:
            if "word/document.xml" not in archive.namelist(): raise ValidationError("DOCX 缺少主文档")
        document = Document(path)
    except (zipfile.BadZipFile, ValueError) as exc:
        raise ValidationError("DOCX 无法打开") from exc
    text = "\n".join(paragraph.text for paragraph in _iter_paragraphs(document))
    if PLACEHOLDER.search(text): raise ValidationError("DOCX 仍存在未替换占位符")
    if snapshot.status == "formal" and INTERNAL_MARKER.search(text): raise ValidationError("DOCX 存在内部来源标记")
    for section in snapshot.snapshot_content.get("sections", []):
        if f"{section['section_code']} {section['title']}" not in text: raise ValidationError(f"DOCX 缺少章节标题：{section['section_code']}")


def _validate_pdf(path: Path) -> None:
    if not path.is_file() or path.stat().st_size < 64 or not path.read_bytes()[:5].startswith(b"%PDF-"):
        raise ValidationError("PDF 文件无效")


def _convert_pdf(docx: Path, workdir: Path) -> Path:
    command = shutil.which("soffice") or shutil.which("libreoffice")
    if not command: raise ValidationError("未检测到 LibreOffice，无法生成 PDF")
    output = workdir / "pdf"; output.mkdir(exist_ok=True)
    profile = workdir / "lo-profile"; profile.mkdir(exist_ok=True)
    result = subprocess.run([command, "-env:UserInstallation=" + profile.as_uri(), "--headless", "--convert-to", "pdf", "--outdir", str(output), str(docx)], check=False, capture_output=True, text=True, timeout=settings.pdf_conversion_timeout)
    pdf = output / f"{docx.stem}.pdf"
    if result.returncode != 0 or not pdf.is_file(): raise ValidationError("PDF 转换失败")
    _validate_pdf(pdf); return pdf


def _persist_artifact(db: Session, job: ReportExportJob, source: Path, fmt: str, snapshot: ReportSnapshot) -> ExportArtifact:
    file_name = safe_filename(f"{snapshot.document_title}_R{snapshot.snapshot_number}", f".{fmt}")
    rel_path = f"exports/{snapshot.document_instance_id}/{job.id}/{uuid.uuid4().hex}.{fmt}"
    target = storage.get_storage().resolve_path(rel_path); target.parent.mkdir(parents=True, exist_ok=True); shutil.copy2(source, target)
    artifact = ExportArtifact(export_job_id=job.id, format=fmt, storage_path=rel_path, file_name=file_name, mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document" if fmt == "docx" else "application/pdf", file_size=target.stat().st_size, sha256=_sha256(target))
    db.add(artifact); db.flush(); return artifact


def start_export(db: Session, snapshot_id: int, report_template_id: int | None, formats: list[str], user: User) -> ReportExportJob:
    snapshot = _owned_snapshot(db, snapshot_id, user)
    template = get_template(db, report_template_id, user) if report_template_id else ensure_default_report_template(db, snapshot.template_id, user)
    from app.services.authorization import require_project_permission
    instance_for_permission = db.get(DocumentInstance, snapshot.document_instance_id)
    require_project_permission(db, user, instance_for_permission.project, "documents.export")
    if sorted(set(formats)) != sorted(formats) or any(item not in {"docx", "pdf"} for item in formats): raise ValidationError("导出格式无效")
    mapped = db.scalar(select(ReportTemplateMapping).where(ReportTemplateMapping.report_template_id == template.id, ReportTemplateMapping.document_template_id == snapshot.template_id))
    if mapped is None: raise ValidationError("Word 模板未绑定该快照的内容模板")
    job = ReportExportJob(report_snapshot_id=snapshot.id, report_template_id=template.id, requested_formats=formats, pdf_status="pending" if "pdf" in formats else "not_requested", started_by=user.id, exporter_version=EXPORTER_VERSION, status="rendering")
    db.add(job); db.flush()
    instance = db.get(DocumentInstance, snapshot.document_instance_id); workflow_service.audit(db, user, "export_started", instance, metadata={"snapshot_id": snapshot.id, "formats": formats})
    workroot = Path(settings.export_temp_dir); workroot.mkdir(parents=True, exist_ok=True)
    try:
        with tempfile.TemporaryDirectory(prefix=f"job-{job.id}-", dir=workroot) as temp:
            directory = Path(temp); docx = directory / "report.docx"
            manifest = _render_docx(snapshot, template, docx); validate_docx(docx, snapshot)
            _persist_artifact(db, job, docx, "docx", snapshot); job.docx_status = "completed"; job.status = "docx_completed"; workflow_service.audit(db, user, "docx_exported", instance, metadata={"snapshot_id": snapshot.id, "job_id": job.id})
            if snapshot.organization_id: tenant_service.record_usage(db, organization_id=snapshot.organization_id, user_id=user.id, project_id=instance.project_id, usage_type="docx_export", quantity=1, unit="export", source_key=f"export_job:{job.id}:docx_export", related_resource_type="export_job", related_resource_id=job.id)
            if "pdf" in formats:
                job.status = "converting_pdf"
                try:
                    pdf = _convert_pdf(docx, directory); _persist_artifact(db, job, pdf, "pdf", snapshot); job.pdf_status = "completed"; workflow_service.audit(db, user, "pdf_exported", instance, metadata={"snapshot_id": snapshot.id, "job_id": job.id})
                    if snapshot.organization_id: tenant_service.record_usage(db, organization_id=snapshot.organization_id, user_id=user.id, project_id=instance.project_id, usage_type="pdf_export", quantity=1, unit="export", source_key=f"export_job:{job.id}:pdf_export", related_resource_type="export_job", related_resource_id=job.id)
                except Exception as exc:
                    job.pdf_status = "failed"; job.status = "partial"; job.error_message = str(exc)[:500]
            job.render_manifest = manifest
            if job.status != "partial": job.status = "completed"
    except Exception as exc:
        job.status = "failed"; job.docx_status = "failed"; job.error_message = str(exc)[:500]
    job.completed_at = datetime.now(timezone.utc); db.commit(); db.refresh(job); return job


def get_export_job(db: Session, job_id: int, user: User) -> ReportExportJob:
    job = db.get(ReportExportJob, job_id)
    if job is not None:
        try: _owned_snapshot(db, job.report_snapshot_id, user)
        except Exception: job = None
    if job is None: raise NotFoundError("导出任务不存在")
    return job


def list_artifacts(db: Session, job_id: int, user: User) -> list[ExportArtifact]:
    get_export_job(db, job_id, user)
    return list(db.scalars(select(ExportArtifact).where(ExportArtifact.export_job_id == job_id).order_by(ExportArtifact.id)))


def get_artifact(db: Session, artifact_id: int, user: User) -> ExportArtifact:
    artifact = db.get(ExportArtifact, artifact_id)
    if artifact is not None:
        try: get_export_job(db, artifact.export_job_id, user)
        except Exception: artifact = None
    if artifact is None: raise NotFoundError("导出文件不存在")
    return artifact


def add_figure(db: Session, instance_id: int, *, project_file_id: int, section_instance_id: int | None, caption: str, sort_order: int, width_inches: float, user: User) -> ReportFigure:
    instance = _owned_instance(db, instance_id, user)
    project_file = db.scalar(select(ProjectFile).where(ProjectFile.id == project_file_id, ProjectFile.project_id == instance.project_id))
    if project_file is None or (project_file.file_type or "").lower() not in {"png", "jpg", "jpeg"}: raise ValidationError("报告图片必须是当前项目的 PNG 或 JPEG 文件")
    if section_instance_id and db.scalar(select(DocumentSectionInstance).where(DocumentSectionInstance.id == section_instance_id, DocumentSectionInstance.document_instance_id == instance.id)) is None: raise ValidationError("图片所属章节不属于当前报告")
    figure = ReportFigure(document_instance_id=instance.id, project_file_id=project_file_id, section_instance_id=section_instance_id, caption=caption, sort_order=sort_order, width_inches=width_inches, created_by=user.id)
    db.add(figure); db.commit(); db.refresh(figure); return figure


def list_figures(db: Session, instance_id: int, user: User) -> list[ReportFigure]:
    instance = _owned_instance(db, instance_id, user)
    return list(db.scalars(select(ReportFigure).where(ReportFigure.document_instance_id == instance.id).order_by(ReportFigure.sort_order, ReportFigure.id)))
