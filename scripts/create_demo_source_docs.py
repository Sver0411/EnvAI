from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# Keep generated demo files portable and independent of the developer's local path.
OUT = Path(__file__).resolve().parents[1] / "demo_assets"
OUT.mkdir(parents=True, exist_ok=True)


def set_font(run, name="STSong", size=10.5, bold=False, color="222222"):
    run.font.name = name
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attribute}"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, *, bold=False, fill=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(str(text))
    set_font(run, bold=bold, color="0B2545" if bold else "222222")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade(cell, fill)


def base_doc(title):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "STSong"
    for attribute in ("ascii", "hAnsi", "eastAsia"):
        normal._element.rPr.rFonts.set(qn(f"w:{attribute}"), "STSong")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color in (("Heading 1", 16, "1F4D78"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 11, "1F4D78")):
        style = doc.styles[name]
        style.font.name = "STSong"
        for attribute in ("ascii", "hAnsi", "eastAsia"):
            style._element.rPr.rFonts.set(qn(f"w:{attribute}"), "STSong")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_font(run, size=20, bold=True, color="0B2545")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("EnvAI 演示项目资料 · 仅用于本地全流程验证")
    set_font(run, size=9, color="666666")
    doc.add_paragraph()
    return doc


def add_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(1.65)
        cells[1].width = Inches(5.5)
        set_cell(cells[0], key, bold=True, fill="E8EEF5")
        set_cell(cells[1], value)
    doc.add_paragraph()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    widths = [6.9 / len(headers)] * len(headers)
    for i, header in enumerate(headers):
        table.columns[i].width = Inches(widths[i])
        set_cell(table.rows[0].cells[i], header, bold=True, fill="D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            set_cell(cells[i], value)
    doc.add_paragraph()


def build_company_profile():
    doc = base_doc("示例环保科技有限公司 环保资料汇编")
    doc.add_heading("一、项目基本信息", level=1)
    add_kv_table(doc, [
        ("项目名称", "示例树脂生产及配套环保设施提升项目"),
        ("建设单位", "示例环保科技有限公司"),
        ("统一社会信用代码", "DEMO-CREDIT-2026-0001"),
        ("项目地址", "示例省示例市环保产业园示例路 18 号"),
        ("行业类别", "C2651 初级形态塑料及合成树脂制造"),
        ("联系人", "示例联系人，400-000-0000"),
        ("占地与建筑", "占地 18,600 m²，建筑面积 12,400 m²"),
    ])
    doc.add_heading("二、生产与污染治理概况", level=1)
    doc.add_paragraph("项目主要生产水性丙烯酸树脂和水性聚氨酯树脂，采用配料、反应、冷却、过滤和灌装工艺。生产车间设置密闭投料和局部收集系统，废气经活性炭吸附浓缩与催化燃烧装置处理后通过 15 m 排气筒排放。")
    doc.add_paragraph("生产废水经厂内污水处理站预处理后接管园区污水处理厂；一般固废分类收集，危险废物暂存于危废库并委托有资质单位处置。")
    doc.add_heading("三、已确认的环保管理要求", level=1)
    for item in [
        "严格执行雨污分流，生产废水不得进入雨水系统。",
        "废气收集系统保持负压运行，活性炭和催化燃烧设施建立运行及更换台账。",
        "危废分类、分区、包装和标识管理，转移过程执行联单制度。",
        "每年组织至少一次突发环境事件综合演练并留存记录。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        for r in p.runs:
            set_font(r)
    path = OUT / "01_企业基本资料汇编.docx"
    doc.save(path)
    return path


def build_risk_inventory():
    doc = base_doc("示例环保科技有限公司 环境风险与应急资源清单")
    doc.add_heading("一、主要环境风险物质", level=1)
    add_table(doc, ["物质", "最大储存量", "储存位置", "主要风险", "防控措施"], [
        ("丙烯酸", "2.0 t", "甲类仓库 A 区", "腐蚀、聚合放热", "冷却、禁火、围堰、泄漏收集"),
        ("异氰酸酯", "1.5 t", "甲类仓库 B 区", "吸入危害、遇水反应", "密闭储存、通风、禁水管理"),
        ("乙酸乙酯", "3.0 t", "甲类仓库 C 区", "易燃、挥发", "防爆电气、静电接地、泡沫灭火"),
        ("天然气", "管道输送", "锅炉房", "易燃、泄漏爆炸", "可燃气体报警、紧急切断"),
    ])
    doc.add_heading("二、主要环保设施与应急资源", level=1)
    add_table(doc, ["设施/资源", "数量或能力", "位置", "责任部门"], [
        ("废气活性炭吸附装置", "2 套，20,000 m³/h", "生产车间屋顶", "设备部"),
        ("催化燃烧装置", "1 套，12,000 m³/h", "废气治理区", "环保部"),
        ("厂内污水处理站", "300 m³/d", "厂区东侧", "环保部"),
        ("消防水池", "600 m³", "厂区西南角", "安环部"),
        ("应急物资柜", "吸附棉、堵漏工具、防护服、呼吸器", "门卫及仓库", "应急办公室"),
    ])
    doc.add_heading("三、应急响应要点", level=1)
    for title, text in [
        ("泄漏事件", "立即停止相关作业，切断物料来源，设置警戒区；由佩戴相应防护装备的人员使用吸附棉和堵漏工具处置，污染物统一作为危险废物收集。"),
        ("废气治理设施故障", "停止相关生产单元，保持故障设施安全状态，通知环保部和设备部，确认修复并完成试运行后方可恢复生产。"),
        ("火灾爆炸事件", "立即报警并启动消防和人员疏散程序，优先切断气源和电源；根据风向和影响范围向园区管委会报告。"),
    ]:
        doc.add_heading(title, level=2)
        doc.add_paragraph(text)
    doc.add_heading("四、资料确认", level=1)
    doc.add_paragraph("以上数据由企业项目组于 2026 年 8 月确认，作为 EnvAI 演示项目的项目事实输入。未在本清单中列出的数量和排放参数不得由模型自行补充。")
    path = OUT / "02_环境风险与应急资源清单.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    for path in (build_company_profile(), build_risk_inventory()):
        print(path)
